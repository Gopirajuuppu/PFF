import re
import os
import pickle
import json
import pandas as pd
from dotenv import load_dotenv
from datetime import datetime
from pprint import pprint
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pydantic import BaseModel, Field
from typing import List

from llama_index.program.openai import OpenAIPydanticProgram
from llama_index.core import ChatPromptTemplate
from llama_index.core.llms import ChatMessage
from llama_index.llms.openai import OpenAI as LlmaOpenAI

from openai import OpenAI

from database_utility import SQLiteDatabase
import prompts
import openai
import sqlite3

from personal_info_extractor import get_personal_details
from DATA.Questions import madrs_rating_guidelines,phq9_rating_guidelines

load_dotenv("creds.env")
os.environ["OPENAI_API_KEY"] = os.getenv('OPENAI_API_KEY')
openai.api_key = os.environ["OPENAI_API_KEY"]

class SummaryDetails(BaseModel):
    """Positive and Negative scenarios from the summary"""

    positive_information: List[str] = Field(default_factory=list, description="Positive information includes any details from the conversation summary that indicate favorable outcomes, improvements, or good practices related to the patient's health. This could encompass positive scenarios such as successful treatments, adherence to medication regimens, positive lifestyle changes, improved symptoms, good habits like regular exercise or healthy eating, and any other positive developments noted during the conversation. It also includes any medications prescribed for managing conditions or promoting wellness.")
    negative_information: List[str] = Field(default_factory=list, description="Negative information comprises any aspects of the conversation summary that highlight adverse outcomes, symptoms, challenges, or concerns regarding the patient's health. This includes negative scenarios such as symptoms experienced by the patient, the presence of diseases or medical conditions, harmful habits like smoking or excessive alcohol consumption, difficulties in managing symptoms or conditions, and any medications prescribed for treating ailments or managing symptoms. Additionally, it covers any lifestyle factors or environmental factors that may negatively impact the patient's health or contribute to the progression of diseases or symptoms.")


class AnalyzerUtil(object):
    def __init__(self, db_name, conversation_table):
        self._db = SQLiteDatabase(db_name)
        self.conversation_table = conversation_table
    
    # Get only completed coversations to analyze
    def get_conversation_id_and_name(self):
        Chat_Type = "chat_completed"
        #data = self._db.fetch_query("SELECT conversation_id, name FROM {} GROUP BY conversation_id, name".format(self.conversation_table))
        query_ = """ SELECT conversation_id, name, Chat_Type FROM {} GROUP BY conversation_id, name, Chat_Type HAVING Chat_Type = 'chat_completed' """.format(self.conversation_table)        
        data = self._db.fetch_query(query_)
        return [conv_id + " | " + name  for conv_id, name, chat_type in data]
    
    def get_conversation(self, conversation_id):
        data = self._db.fetch_data(self.conversation_table, condition="conversation_id='{}' ".format(conversation_id))           
        conversation_str = ""
        for row in data:
            if row[3] == "user":
                conversation_str += "<span style='color:#00FF00'>{}</span> : {}<hr style='border:0.25px solid white'>".format(row[3], re.sub(r'\n+', '<br>', row[4]).strip())
            else:
                conversation_str += "<span style='color:#FF00FF'>{}</span> : {}<hr style='border:0.25px solid white'>".format(row[3], re.sub(r'\n+', '<br>', row[4]).strip())
        return conversation_str
    
    # Updating ChatType Column if there are any completed conversation in the entire conversation
    def update_chatType(self):        
        chat_complete_update_query = """ UPDATE {}
                            SET Chat_Type = 'chat_completed'
                            WHERE conversation_id IN (
                            SELECT conversation_id
                            FROM {}
                            WHERE Chat_Type = 'chat_completed'
                            )
                            """.format(self.conversation_table,self.conversation_table)
        self._db.update_query(chat_complete_update_query)
    
    # get the audio data from DB as per the selection of the conversation ID
    def get_audio_data(self,selected_conv_id):
        query_to_get_audio_data = """ SELECT conversation_id, audio FROM conversation_log_table where conversation_id='{}' """.format(selected_conv_id)        
        audio_data = self._db.fetch_query(query_to_get_audio_data)  
        # Iterate through the rows and save each audio file
        for index, row in enumerate(audio_data):
            binary_audio = row[1] 
            audio_directory = './DATA/Audio_Files/'+selected_conv_id+'_audio_{}.wav'.format(index)
            # Save the audio bytes to a temporary WAV file
            with open(audio_directory, "wb") as audio_file:        
                audio_file.write(binary_audio)
                audio_file.close()
            print(f'Saved {audio_directory}')
        
class ConversationAnalyzer(object):
    #def __init__(self, db_name, conversation_table, analyzer_table, questionnaire_json_file):
    def __init__(self, db_name, conversation_table, analyzer_table):
        self._client = OpenAI()
        self._db = SQLiteDatabase(db_name)
        self._db_name = db_name
        self._llama_llm = LlmaOpenAI(model="gpt-3.5-turbo-0613", temperature=0.2)
        
        self._conversation_table = conversation_table
        self._analyzer_table = analyzer_table
        
        # Read promis questions
        # with open(questionnaire_json_file) as jobj:
        #     self._question_data = json.load(jobj)
            
        
    def get_summary(self, section, conversation):
        analysis_prompt = prompts.GET_SUMMARY_ANALYSIS_PROMPT
        pos_neg_prompt = ("Extract both positive and negative information from the below `Summary Content` according "
                          "to provided JSON schema. EXTRACT EXACT SHORT PHRASES WITHOUT ANY MODIFICATION.\n"
                          "\n"
                          "## Summary Content\n"
                          "{summary_content}")
        analysis_prompt = analysis_prompt.format(section, conversation)
        
        messages = [{"role": "system", "content": analysis_prompt}]
        response = self._client.chat.completions.create(
                                                                  model="gpt-3.5-turbo-1106", #gpt-4-1106-preview, gpt-3.5-turbo-1106
                                                                  temperature=0.5,
                                                                  messages=messages
                                                                )
        summary_content = response.choices[0].message.content
        print("summary_content", summary_content)
        program = OpenAIPydanticProgram.from_defaults(
                    output_cls=SummaryDetails,
                    llm=self._llama_llm,
                    prompt_template_str=pos_neg_prompt,
                    verbose=False,
                )
        pos_neg_output = program(summary_content=summary_content)
        print("pos_neg_output", pos_neg_output.model_dump())
        positive_phrases = '|'.join(re.escape(phrase) for phrase in pos_neg_output.model_dump()['positive_information'])
        negative_phrases = '|'.join(re.escape(phrase) for phrase in pos_neg_output.model_dump()['negative_information'])
        if positive_phrases:
            summary_content = re.sub(re.compile(r'\b(' + positive_phrases + r')\b', re.IGNORECASE), r'<GREEN>\1</GREEN>', summary_content)
        if negative_phrases:
            summary_content = re.sub(re.compile(r'\b(' + negative_phrases + r')\b', re.IGNORECASE), r'<RED>\1</RED>', summary_content)        
        return summary_content.replace("<RED>", "<span style='color:#FF0000'>").replace("</RED>", "</span>").replace("<GREEN>", "<span style='color:#00FF00'>").replace("</GREEN>", "</span>").replace("\n", "<br>")

    
    
    def get_choices(self, section, conversation, choices, questions, choices_dict):
        analysis_prompt = prompts.GET_CHOICES_ANALYSIS_PROMPT
        analysis_prompt = analysis_prompt.format(section, choices, questions, conversation, choices_dict)
        
        messages = [{"role": "system", "content": analysis_prompt}]
        response = self._client.chat.completions.create(
                                                                  model="gpt-3.5-turbo-1106",
                                                                  temperature=0.2,
                                                                  messages=messages
                                                                )
        response = response.choices[0].message.content            
        response_dict = {int(k): v for k, v in json.loads(response).items()} 
        return response_dict
    
    def get_section_ratings_choices(self,section, conversation, choices, questions, choices_dict):
        analysis_prompt = ""        
        if self.mcq_type == 'MADRS':            
            mcq_prompt = prompts.CHOICE_ANALYSIS_PROMPT         
            SYMTOMS_TO_IDENTIFY,MEDICAL_KEY_POINTS = madrs_rating_guidelines.SYMTOMS_TO_IDENTIFY,madrs_rating_guidelines.MEDICAL_KEY_POINTS
            guide_lines = madrs_rating_guidelines.MADRS_RATING_GUIDE_LINES
            rating_guide_lines=guide_lines[section]['rating_guide_lines']
            example_conversation=guide_lines[section]['example_conversation']
            example_rating=guide_lines[section]['example_rating']
            example_symptems_key_points=guide_lines[section]['example_Symptems_key_points']            
            analysis_prompt = mcq_prompt.format(section, conversation, questions,
                                                 SYMTOMS_TO_IDENTIFY, 
                                                 MEDICAL_KEY_POINTS,
                                                 rating_guide_lines,
                                                 example_conversation,
                                                 example_symptems_key_points,
                                                 example_rating,
                                                 choices                                                 
                                                )

        elif self.mcq_type == 'PHQ9':            
            mcq_prompt = prompts.CHOICE_ANALYSIS_PROMPT         
            SYMTOMS_TO_IDENTIFY,MEDICAL_KEY_POINTS = phq9_rating_guidelines.SYMTOMS_TO_IDENTIFY,phq9_rating_guidelines.MEDICAL_KEY_POINTS
            guide_lines = phq9_rating_guidelines.PHQ9_RATING_GUIDE_LINES
            rating_guide_lines=guide_lines[section]['rating_guide_lines']
            example_conversation=guide_lines[section]['example_conversation']
            example_rating=guide_lines[section]['example_rating']
            example_symptems_key_points=guide_lines[section]['example_Symptems_key_points']            
            analysis_prompt = mcq_prompt.format(section, conversation, questions,
                                                 SYMTOMS_TO_IDENTIFY, 
                                                 MEDICAL_KEY_POINTS,
                                                 rating_guide_lines,
                                                 example_conversation,
                                                 example_symptems_key_points,
                                                 example_rating,
                                                 choices                                                 
                                                )  
        messages = [{"role": "system", "content": analysis_prompt}]
        response = self._client.chat.completions.create(
                                                  model="gpt-4",
                                                  temperature=0.2,
                                                  messages=messages
                                                    )
        response = response.choices[0].message.content 
        #response_dict = {k: v for k, v in json.loads(response).items()}
        return response
    
    def map_total_score(self,mcq_type,total_score):
        section_rating_guideLines=""
        if mcq_type=="MADRS":
            section_rating_guideLines = madrs_rating_guidelines.FINAL_RATING_GUIDELINES
        elif mcq_type=="PHQ9":
            section_rating_guideLines = phq9_rating_guidelines.FINAL_RATING_GUIDELINES
        for score_range, description in section_rating_guideLines.items():
            start, end = map(int, score_range.split('-')) if '-' in score_range else (int(score_range), int(score_range))
            if start <= total_score <= end:
                return description
        return "an undefined level of symptoms"
    
    def get_final_summary(self,mcq_type,choice_summary):
        summary = []
        total_score = 0 
        for section, details in choice_summary.items():
            section_rating = int(details['section_rating'])
            summary.append(f"{section_rating} on {section}")
            total_score += section_rating
        summary_str = ", ".join(summary)           
        # Mapping the total score to the appropriate description
        total_score_description = self.map_total_score(mcq_type,total_score) 
        final_summary = f"""Based on the conversation with the patient, the patient scores {summary_str}.
        The total {mcq_type} score is {total_score}. This indicates that the patient has {total_score_description}.
        """
        print("Final Summary is ----- ",final_summary)
        return final_summary    
    
    def generate_docx(self, all_section_questions_and_choices, analysis_summary, conversation_id,personal_info_df,final_summary):
        #docx_output = "questionnaire.docx"
        docx_output = "./DATA/Outputs/{}_questionnaire.docx".format(str(conversation_id))
        doc = Document()
        # Add heading
        heading = doc.add_heading("Pre-Medical Form", level=1)
        doc.add_paragraph().add_run().add_break()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].underline = True
        heading.runs[0].bold = True
        heading.runs[0].font.color.rgb = RGBColor(0x00, 0x4B, 0x87)
        heading.runs[0].font.name = "HCLTech Roobert"

        for section, section_details in all_section_questions_and_choices.items():
            section_para = doc.add_paragraph(section)
            section_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            section_para.runs[0].font.color.rgb = RGBColor(0x22, 0x8b, 0x22)
            section_para.runs[0].font.name = "HCLTech Roobert"
            section_para.runs[0].bold = True

            for i, question in enumerate(section_details["questions"]):
                question_para = doc.add_paragraph(question)
                question_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                question_para.runs[0].font.color.rgb = RGBColor(0x00, 0x4B, 0x87)
                question_para.runs[0].font.name = "HCLTech Roobert"
                question_para.runs[0].bold = True

                choises_str = ""
                for choice in section_details["all_choices"]:
                    if choice == section_details["selected_choices"][i]:
                        choises_str = choises_str + f"☑ {choice}" + "  "
                    else:
                        choises_str = choises_str + f"☐ {choice}" + "  "
                choice_para = doc.add_paragraph(choises_str)
                choice_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                choice_para.runs[0].font.color.rgb = RGBColor(0x00, 0x4B, 0x87)
                choice_para.runs[0].font.name = "HCLTech Roobert"
        doc.add_paragraph().add_run().add_break()
        
        # Adding Personal Details
        if len(personal_info_df)!=0:
            section_para = doc.add_paragraph("Personal Details")
            section_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            section_para.runs[0].font.color.rgb = RGBColor(0x22, 0x8b, 0x22)
            section_para.runs[0].font.name = "HCLTech Roobert"
            section_para.runs[0].bold = True
                
            personal_info_df = personal_info_df.reset_index()              
            t = doc.add_table(rows=personal_info_df.shape[0], cols=personal_info_df.shape[1])        
            t.style = 'TableGrid'        
            # Add the column headings
            for j in range(personal_info_df.shape[1]):
                t.cell(0, j).text = personal_info_df.columns[j]            
            # Add the body of the data frame to the table
            for i in range(personal_info_df.shape[0]):
                for j in range(personal_info_df.shape[1]):
                    cell = personal_info_df.iat[i, j]
                    t.cell(i, j).text = str(cell) 
            doc.add_paragraph().add_run().add_break()

        for section, summary in analysis_summary.items():
            if section != "PERSONAL_INFORMATION":
                summary = summary.replace("<span style='color:#FF0000'>", "<RED>")\
                                 .replace("</span>", "</RED>")\
                                 .replace("<span style='color:#00FF00'>", "<GREEN>")\
                                 .replace("</span>", "</GREEN>")\
                                 .replace("<br>", "\n")
                section_para = doc.add_paragraph(section)
                section_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                section_para.runs[0].font.color.rgb = RGBColor(0x22, 0x8b, 0x22)
                section_para.runs[0].font.name = "HCLTech Roobert"
                section_para.runs[0].bold = True

                summary_paragraph = doc.add_paragraph()
                summary_runs = summary.split("<")

                for run in summary_runs:
                    if ">" in run:
                        tag, content = run.split(">")
                        run = summary_paragraph.add_run(content)
                        if tag.startswith("RED"):
                            run.font.color.rgb = RGBColor(255, 0, 0)  # RED
                        elif tag.startswith("GREEN"):
                            run.font.color.rgb = RGBColor(0, 128, 0)  # GREEN
                    else:
                        summary_paragraph.add_run(run)
        # Adding Final Summary
        section_para = doc.add_paragraph("Final Choice Summary")
        section_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        section_para.runs[0].font.color.rgb = RGBColor(0x22, 0x8b, 0x22)
        section_para.runs[0].font.name = "HCLTech Roobert"
        section_para.runs[0].bold = True
        doc.add_paragraph(final_summary)

        doc.save(docx_output)
        return docx_output
    
    
    def insert_analyzed_conversation(self, analysis_summary, choice_summary, conversation_id_option,personal_info_df,final_summary):
        conversation_id = conversation_id_option.split("|")[0].strip()
        conversation_name = conversation_id_option.split("|")[1].strip()
        all_section_questions_and_choices = dict()
        
        serialized_analysis_summary = pickle.dumps(analysis_summary)
        serialized_raw_choices_summary = pickle.dumps(choice_summary)
        
        for idx1, (section, report) in enumerate(choice_summary.items()):
            questions = list(report['choice_summary'].keys())
            selected_choices = list(report['choice_summary'].values())
            all_choices = report['all_choices']
            all_section_questions_and_choices[section] = {"questions": questions, "selected_choices": selected_choices, "all_choices": all_choices}
            
        serialized_master_dict = pickle.dumps(all_section_questions_and_choices)
        
        # Write to DOCX
        output_docx = self.generate_docx(all_section_questions_and_choices, analysis_summary, conversation_id,personal_info_df,final_summary)
        with open(output_docx, 'rb') as file_object:
            serialized_docx = file_object.read()
            
        os.remove(output_docx) if os.path.exists(output_docx) else None
        current_datetime = datetime.now().strftime("%d-%m-%Y %H:%M:%S")
        self._db.insert_single_row(self._analyzer_table,
                                                  [conversation_id,
                                                   conversation_name,
                                                   current_datetime,
                                                   serialized_analysis_summary,
                                                   serialized_raw_choices_summary,
                                                   serialized_master_dict,
                                                   serialized_docx])
        
        
    def get_analyzed_conversation(self, conversation_id_option):
        conversation_id = conversation_id_option.split("|")[0].strip()
        data = self._db.fetch_data(self._analyzer_table, condition="conversation_id='{}' ".format(conversation_id)) 
        if not data:
            return False, None, None
        else:
            serialized_analysis_summary = data[0][3]
            serialized_raw_choices_summary = data[0][4]            
            analysis_summary = pickle.loads(serialized_analysis_summary)
            raw_choices_summary = pickle.loads(serialized_raw_choices_summary)
            return True, analysis_summary, raw_choices_summary
        
        
    
    def __call__(self, conversation_id_option):
        self._conversation_id = conversation_id_option.split("|")[0].strip()
        
        self.sections = list()
        self.analysis_summary = dict()
        self.choice_summary = dict()
        self.personal_info_df = pd.DataFrame()
        
        self.mcq_type = self._db.fetch_query(""" select Mcq_Type from {} WHERE conversation_id='{}' """.format(self._conversation_table,self._conversation_id))[0][0]
        questionary_json_file=''
        if self.mcq_type=='PHQ9':
            questionary_json_file = './DATA/Questions/phq9.json'
        elif self.mcq_type=='MADRS':
            questionary_json_file = './DATA/Questions/madrs.json'
        elif self.mcq_type=='PROMIS-29':
            questionary_json_file = './DATA/Questions/promis.json'
        with open(questionary_json_file) as jobj:
            self._question_data = json.load(jobj)
        
        # Get conversation sections
        #section_data = self._db.fetch_query("SELECT section FROM {} WHERE conversation_id='{}' GROUP BY section".format(self._conversation_table, self._conversation_id))
        self.sections =[list(sections_data.keys())[0] for sections_data in self._question_data]       
#         for row in section_data:
#             self.sections.append(row[0])     
           
        for index, section in enumerate(self.sections): 
            print("Processing for the section...",section)           
            conversation_raw_str = ""
            # Fetch section wise conversation data for the specified conversation id
            section_conv = self._db.fetch_data(self._conversation_table, condition="conversation_id='{}' and section='{}' ".format(self._conversation_id, section))
            for row in section_conv:                
                if row[3] == "user":
                    conversation_raw_str += "{} : {}\n".format("Patient", re.sub(r'\n+', '\n', row[4]).strip())
                else:
                    conversation_raw_str += "{} : {}\n".format("Doctor", re.sub(r'\n+', '\n', row[4]).strip())
            choices = "- " + "\n- ".join(self._question_data[index][section]["choices"])
            choices_dict = dict(enumerate(self._question_data[index][section]["choices"], 1))
            questions = "\n".join([str(idx+1) + ". " + qtn for idx, qtn in enumerate(self._question_data[index][section]["questions"])])
            # Get summary
            summary = self.get_summary(section, conversation_raw_str)
            
            if section =="PERSONAL_INFORMATION":
                user_info = get_personal_details(summary)                                               
                user_info_df = pd.DataFrame.from_dict(data=user_info,orient='index',columns=['Personal Info'])
                self.personal_info_df = user_info_df    
            else: 
                #predicted_choices = self.get_choices(section, conversation_raw_str, choices, questions, choices_dict)
                predicted_ratings_choices = self.get_section_ratings_choices(section, conversation_raw_str, choices, questions, choices_dict)
                predicted_choices = json.loads(predicted_ratings_choices)
                section_rating = predicted_choices['section_rating']['rating']
                choice_summary_dict = {k: v['choice'] for k, v in predicted_choices.items()} 
                
            # Capture analysis
            self.analysis_summary[section] = summary
            if self._question_data[index][section]["choices"]:
                questions_list = self._question_data[index][section]["questions"]
                questions_list.append(f"""Overall Symptoms Identified-{section}""")                
                self.choice_summary[section] = {"choice_summary": {question: choice_summary_dict[key] for key, question in zip(choice_summary_dict.keys(), questions_list)},
                                                "all_choices": self._question_data[index][section]["choices"],
                                                "section_rating":section_rating}
            
        #Adding Final Summary for all the sections
        final_summary = self.get_final_summary(self.mcq_type, self.choice_summary)
        
        # Insert to db  
        self.insert_analyzed_conversation(self.analysis_summary, self.choice_summary, conversation_id_option,self.personal_info_df,final_summary)

        return self.analysis_summary, self.choice_summary,self.personal_info_df,final_summary
            
        